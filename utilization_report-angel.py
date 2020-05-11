#!/usr/bin/python
from selenium import webdriver
from datetime import date
import os
from docx2pdf import convert
from docx import Document
from docx.shared import Inches
import owncloud
#from io import StringIO
today = date.today()
D = today.strftime("%d")
print("D =", D)
M = today.strftime("%m")
print("M =", M)
Y = today.strftime("%Y")
print("Y =", Y)
if not os.path.exists(f'/Users/mani/Documents/Angel-Monring/Screenshot-{D}-{M}-{Y}'):
    os.makedirs(f'/Users/mani/Documents/Angel-Morning/Screenshot-{D}-{M}-{Y}')
driver=webdriver.Chrome(executable_path="/Users/mani/bin/chromedriver")
SP=f"/Users/mani/Documents/Angel-Morning/Screenshot-{D}-{M}-{Y}"
driver.set_window_size(620, 510) # set the window size that you need 
url=[]
#WEB1
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-WEB-1&srv=SYS-CPU-Usage&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-WEB-1&srv=SYS-Memory-Usage&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-WEB-1&srv=SYS-Current-Load&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-WEB-1&srv=MSF-Conn-WS-443&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
#WEB2
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-WEB-2&srv=SYS-CPU-Usage&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-WEB-2&srv=SYS-Memory-Usage&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-WEB-2&srv=SYS-Current-Load&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-WEB-2&srv=MSF-Conn-WS-443&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
#STREAMER1
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-1&srv=SYS-CPU-Usage&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-1&srv=SYS-Memory-Usage&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-1&srv=SYS-Current-Load&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-1&srv=MSF-Conn-WS-Pushserver-8443&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
#STREAMER-2
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-2&srv=SYS-CPU-Usage&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-2&srv=SYS-Memory-Usage&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-2&srv=SYS-Current-Load&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-2&srv=MSF-Conn-WS-Pushserver-8443&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
#STREAMER-3
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-3&srv=SYS-CPU-Usage&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-3&srv=SYS-Memory-Usage&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-3&srv=SYS-Current-Load&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-3&srv=MSF-Conn-WS-Pushserver-8443&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
#STREAMER-4
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-4&srv=SYS-CPU-Usage&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-4&srv=SYS-MEMORY-Usage&start=05%2F07%2F2020+09%3A00+&end=05%2F07%2F2020+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-4&srv=SYS-Current-Load&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-4&srv=MSF-Conn-WS-Pushserver-8443&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
#STREAMER-5
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-5&srv=SYS-CPU-Usage&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-4&srv=SYS-MEMORY-Usage&start=05%2F07%2F2020+09%3A00+&end=05%2F07%2F2020+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-5&srv=SYS-Current-Load&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-5&srv=MSF-Conn-WS-Pushserver-8443&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
#STREAMER-6
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-6&srv=SYS-CPU-Usage&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-4&srv=SYS-MEMORY-Usage&start=05%2F07%2F2020+09%3A00+&end=05%2F07%2F2020+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-6&srv=SYS-Current-Load&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
url.append(f'https://Abmaproductteam:noc@123@mprod-ws.angelbroking.com/pnp4nagios/graph?host=Angel-Tr-Prod-BINARY-STREAMER-6&srv=MSF-Conn-WS-Pushserver-8443&start={M}%2F{D}%2F{Y}+09%3A00+&end={M}%2F{D}%2F{Y}+10%3A00+')
i=1
for x in url:
    driver.get(x)
    driver.set_page_load_timeout(20)
    filename = 'morning'+str(i)+'.png'
    i+=1
    print ({filename})
    
    driver.save_screenshot(f'{SP}/{filename}')
driver.close()
document = Document()
p = document.add_paragraph()
r = p.add_run()
for i in range(1,32):
 filename = 'morning'+str(i)+'.png'
 #print ({filename})
 r.add_picture(f'{SP}/{filename}', width=Inches(6.7))
document.save(f'{SP}/Angel-TR-Utilization-Report-{D}{M}{Y}9to10AM.docx')
convert(f'{SP}/Angel-TR-Utilization-Report-{D}{M}{Y}9to10AM.docx', f'{SP}/Angel-TR-Utilization-Report-{D}{M}{Y}9to10AM.pdf')
oc = owncloud.Client('https://files.marketsimplified.com/owncloud/')
oc.login('infraadmin', 'File$Upl0ad')
oc.mkdir(f'INFRA/Utilisation-Reports/ANGEL/OMNESYS/{D}{M}{Y}/')
oc.mkdir(f'INFRA/Utilisation-Reports/ANGEL/OMNESYS/{D}{M}{Y}/Morning/')
oc.put_file(f'INFRA/Utilisation-Reports/ANGEL/OMNESYS/{D}{M}{Y}/Morning/', f'{SP}/Angel-TR-Utilization-Report-{D}{M}{Y}9to10AM.pdf')
